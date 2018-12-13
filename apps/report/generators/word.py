from io import BytesIO

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def makedocx(data: list) -> bytes:
    # Создаем документ docx
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    paragraph = document.add_paragraph()
    p_format = paragraph.paragraph_format
    p_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.add_run('Результат поиска информации в ИПС “Флоризель”')
    font = run.font
    font.bold = True
    first_paragraph = True
    for item in data:
        phone = item["target"]
        current_paragraph = document.add_paragraph()
        prefix = '' if first_paragraph else '\n'
        current_paragraph.add_run(prefix + 'Поиск по ')
        run = current_paragraph.add_run(phone)
        font = run.font
        font.bold = True
        run = current_paragraph.add_run(', дал следующий результат:')
        font = run.font
        font.bold = False
        prefix = ''
        for result in item["results"]:
            # Добавляем Источник перед таблицей
            db_name = result["index_name"]
            document.add_paragraph().add_run(prefix + 'источник - база данных «' + db_name + '»')
            cnt = 1
            for attributes in result["data"]:
                table = document.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                prefix = '\n'
                for kvp in attributes.items():
                    try:
                        if kvp[1] != "" and kvp[0] not in ["db_shard", "id_identity", "ID_Obj"]:
                            # print(kvp)
                            row_cells = table.add_row().cells
                            row_cells[0].text = kvp[0]
                            if kvp[1] is None:
                                row_cells[1].text = ''
                            else:
                                row_cells[1].text = str(kvp[1])
                    except ValueError:
                        continue
                if cnt < len(result["data"]):
                    document.add_paragraph().add_run("")
                    cnt += 1
        document.add_paragraph().add_run("")
    buf = BytesIO()
    document.save(buf)
    res = buf.getvalue()
    buf.close()
    return res
